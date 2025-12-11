import csv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from collections import OrderedDict

def add_table_borders(table):
    """Add borders to table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def format_table_name(table_name):
    """
    Convert table name to a readable heading
    Examples:
      demographics_part_1 → Demographics Part 1
      shoulder_diagnosis → Shoulder Diagnosis
      general_assessed_rom → General Assessed ROM
    """
    # Replace underscores with spaces and title case
    formatted = table_name.replace('_', ' ').title()
    
    # Handle special cases
    formatted = formatted.replace('Rom', 'ROM')
    formatted = formatted.replace('Pruj', 'PRUJ')
    formatted = formatted.replace('Druj', 'DRUJ')
    formatted = formatted.replace('Mucl', 'MUCL')
    formatted = formatted.replace('Mepi', 'MEPI')
    formatted = formatted.replace('Pree', 'PREE')
    formatted = formatted.replace('Ases', 'ASES')
    formatted = formatted.replace('Vas', 'VAS')
    
    return formatted


def create_carbone_template_from_csv(csv_file_path, output_docx_path):
    """
    Create a Carbone template DOCX organized by SQL table names.
    
    Structure:
      - Heading: Table Name (formatted)
      - Table with 2 columns:
        - Column 1: Label (from CSV)
        - Column 2: {d.sql_column_name} (Carbone placeholder)
    
    Args:
        csv_file_path: Path to column_mapping.csv
        output_docx_path: Path for output DOCX template
    """
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Clinical Documentation Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add encounter ID
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('Encounter ID: {d.encounter_id}')
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(0, 0, 139)
    subtitle_run.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Read CSV and organize by table
    tables_data = OrderedDict()
    
    try:
        with open(csv_file_path, 'r', encoding='utf-8') as csvfile:
            csv_reader = csv.DictReader(csvfile)
            
            for row in csv_reader:
                label = row['Label']
                sql_column = row['SQL_Column_Name']
                table_name = row['Table_Name']
                category = row['Category']
                
                # Skip empty rows
                if not label or not sql_column:
                    continue
                
                # Skip system fields
                if sql_column in ['id', 'created_at', 'updated_at']:
                    continue
                
                # Initialize table if not exists
                if table_name not in tables_data:
                    tables_data[table_name] = {
                        'category': category,
                        'fields': []
                    }
                
                # Add field to table
                tables_data[table_name]['fields'].append({
                    'label': label,
                    'sql_column': sql_column
                })
        
        total_tables = len(tables_data)
        total_fields = sum(len(table['fields']) for table in tables_data.values())
        
        print(f"✅ Loaded {total_fields} fields from {total_tables} tables")
        
    except FileNotFoundError:
        print(f"❌ Error: File '{csv_file_path}' not found")
        return
    except Exception as e:
        print(f"❌ Error reading CSV file: {e}")
        return
    
    # Track categories for organization
    current_category = None
    table_count = 0
    
    # Generate document sections organized by table
    for table_name, table_info in tables_data.items():
        
        table_count += 1
        category = table_info['category']
        fields = table_info['fields']
        
        # Skip tables with no fields
        if len(fields) == 0:
            continue
        
        # Add category heading when category changes
        if category != current_category:
            if current_category is not None:
                doc.add_page_break()
            
            category_heading = doc.add_heading(category.upper(), 1)
            category_heading.runs[0].font.color.rgb = RGBColor(0, 0, 139)
            current_category = category
        
        # Add table name as section heading (Level 2)
        formatted_table_name = format_table_name(table_name)
        section_heading = doc.add_heading(formatted_table_name, 2)
        section_heading.runs[0].font.color.rgb = RGBColor(70, 70, 70)
        
        # Create data table with label and value columns
        data_table = doc.add_table(rows=0, cols=2)
        data_table.style = 'Light Grid Accent 1'
        
        # Set column widths
        data_table.columns[0].width = Inches(3.5)  # Label column
        data_table.columns[1].width = Inches(3.0)  # Value column
        
        # Add each field as a row
        for field in fields:
            row = data_table.add_row()
            
            # Column 1: Label
            label_text = field['label']
            # Clean up label (remove HTML tags)
            label_text = label_text.replace('<', '').replace('>', '')
            if len(label_text) > 100:
                label_text = label_text[:97] + '...'
            
            row.cells[0].text = label_text
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            
            # Column 2: Carbone placeholder
            row.cells[1].text = f"{{d.{field['sql_column']}}}"
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
            row.cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)
        
        # Add borders to table
        add_table_borders(data_table)
        
        # Add spacing after table
        doc.add_paragraph()
        
        # Print progress
        if table_count % 10 == 0:
            print(f"  Processed {table_count}/{total_tables} tables...")
    
    # Add usage instructions at the end
    doc.add_page_break()
    doc.add_heading('Template Usage Instructions', 1)
    
    instructions = doc.add_paragraph()
    instructions.add_run('About This Carbone Template:\n\n').bold = True
    
    inst_text = instructions.add_run(
        '• This document is organized by database tables (shown as Level 2 headings)\n'
        '• Each table contains fields with:\n'
        '  - Left column: Human-readable label\n'
        '  - Right column: {d.column_name} placeholder for Carbone\n\n'
        '• When processed by Carbone, placeholders are replaced with actual data\n'
        '• The encounter_id links all data for a single patient visit\n\n'
        'Customization:\n'
        '• Delete entire tables/sections you don\'t need\n'
        '• Rearrange fields in any order\n'
        '• Copy fields to create custom layouts\n'
        '• Add text, formatting, images, and conditional logic\n'
    )
    inst_text.font.size = Pt(10)
    
    # Save document
    doc.save(output_docx_path)
    
    print(f"\n{'='*70}")
    print(f"✅ Carbone template created successfully!")
    print(f"📄 Output: {output_docx_path}")
    print(f"📊 Total tables: {total_tables}")
    print(f"📋 Total fields: {total_fields}")
    print(f"{'='*70}\n")


def create_simplified_template(csv_file_path, output_docx_path):
    """
    Create a simplified template with only key demographics fields
    Good for testing or minimal reports.
    """
    doc = Document()
    
    # Title
    title = doc.add_heading('Patient Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Encounter ID
    para = doc.add_paragraph()
    run = para.add_run('Encounter: {d.encounter_id}')
    run.bold = True
    run.font.size = Pt(12)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Read CSV to find key fields
    key_fields = []
    
    with open(csv_file_path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        
        # Fields to include in simplified template
        priority_originals = [
            'record_id', 'mrn', 'sex', 'lastname', 'givenname',
            'dateofbirth', 'email_pt', 'mobile', 'date', 
            'body_region', 'insurance_type'
        ]
        
        for row in reader:
            orig = row['Original_Field_ID'].lower()
            
            for priority in priority_originals:
                if priority in orig and row['Table_Name'] == 'demographics_part_1':
                    key_fields.append({
                        'label': row['Label'],
                        'sql_column': row['SQL_Column_Name']
                    })
                    break
    
    # Create simple table
    doc.add_heading('Patient Information', 1)
    
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light List Accent 1'
    
    for field in key_fields:
        row = table.add_row()
        row.cells[0].text = field['label']
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        
        row.cells[1].text = f"{{d.{field['sql_column']}}}"
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].paragraphs[0].runs[0].font.name = 'Courier New'
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 0, 255)
    
    doc.save(output_docx_path)
    print(f"✅ Simplified template created: {output_docx_path}")


# Main execution
if __name__ == "__main__":
    csv_input = "/Users/ankitgatg/Desktop/RedCAP-Terraform/ansible/files/sql_schema/5/column_mapping.csv"
    
    print("="*70)
    print("CARBONE TEMPLATE GENERATOR")
    print("="*70)
    print(f"Input: {csv_input}")
    print("="*70)
    print()
    
    # Generate full comprehensive template
    print("1️⃣  Creating comprehensive template (all tables)...")
    full_output = "patient_report.docx"
    create_carbone_template_from_csv(csv_input, full_output)
    
    # Generate simplified template
    print("\n2️⃣  Creating simplified template (key fields only)...")
    simple_output = "patient_report_simple.docx"
    create_simplified_template(csv_input, simple_output)
    
    print("\n✅ Both templates created successfully!")
    print("\nFiles generated:")
    print(f"  📄 {full_output}")
    print(f"     → Comprehensive template with all {2662} fields")
    print(f"  📄 {simple_output}")
    print(f"     → Simplified template with ~12 key demographics fields")
    print("\n🎯 Upload to S3 template bucket and test with Carbone Lambda")